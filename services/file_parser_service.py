"""File parser service: extract plain text from uploaded documents.

Supports .txt, .pdf, .docx uploads. Returns extracted text content.
"""

import io
import logging
import os

logger = logging.getLogger(__name__)

# Maximum file size: 50 MB
MAX_FILE_SIZE = 50 * 1024 * 1024

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
ALLOWED_MIMETYPES = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class ParseError(Exception):
    """Raised when file parsing fails."""


def validate_file(filename: str, file_size: int) -> str:
    """Validate file extension and size. Returns normalised extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ParseError(
            f"Unsupported file type '{ext}'. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    if file_size > MAX_FILE_SIZE:
        raise ParseError(
            f"File too large ({file_size / 1024 / 1024:.1f} MB). Maximum is {MAX_FILE_SIZE / 1024 / 1024:.0f} MB."
        )
    return ext


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from file bytes based on extension.

    Returns extracted text content. Raises ParseError on failure.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return _parse_txt(file_bytes)
    elif ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    else:
        raise ParseError(f"Unsupported file type: {ext}")


def _parse_txt(file_bytes: bytes) -> str:
    """Parse plain text file. Try UTF-8, fall back to latin-1."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ParseError("PDF support requires PyPDF2. Install with: pip install PyPDF2")

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        content = "\n\n".join(pages)
        if not content.strip():
            raise ParseError("PDF appears to contain no extractable text (may be scanned/image-based).")
        return content
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"Failed to parse PDF: {e}")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx.

    Extracts from paragraphs, tables, headers, and footers to handle
    documents that store content in any of these structures.
    """
    try:
        from docx import Document
    except ImportError:
        raise ParseError("DOCX support requires python-docx. Install with: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # Body paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Tables — extract each cell's text row by row
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # Headers and footers
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer and header_footer.is_linked_to_previous is False:
                    for p in header_footer.paragraphs:
                        text = p.text.strip()
                        if text:
                            parts.append(text)

        content = "\n\n".join(parts)
        if not content.strip():
            raise ParseError("DOCX appears to contain no text content.")
        return content
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"Failed to parse DOCX: {e}")
